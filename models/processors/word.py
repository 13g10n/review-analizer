import settings

from docx import Document


class WordProcessor:

    @staticmethod
    def generate(summary):
        document = Document()

        for report in summary.reports:
            document.add_heading('Text report for "{}"'.format(report.excerpt), 1)

            document.add_heading('Input', 2)
            document.add_paragraph('{}'.format(report.content))

            document.add_heading('Summary', 2)
            summary_table = document.add_table(rows=2, cols=3)
            summary_table.style = 'Table Grid'
            WordProcessor.__add_summary_table_rows(summary_table, report)

            document.add_heading('Details', 2)
            details_table = document.add_table(rows=0, cols=7)
            details_table.style = 'Table Grid'
            WordProcessor.__add_sentences_reports_heading_row(details_table)
            for sentence_report in report.reports:
                first_row = WordProcessor.__add_sentence_report_row(details_table, sentence_report)
                WordProcessor.__add_keyword_reports_heading_row(details_table)
                last_row = first_row
                for keyword_report in sentence_report.reports:
                    last_row = WordProcessor.__add_keyword_report_row(details_table, keyword_report)
                first_row.cells[0].merge(last_row.cells[0])

        return document

    @staticmethod
    def __add_summary_table_rows(table, report):
        for index, value in enumerate(["Sentences", "Evaluation", "Summary"]):
            table.rows[0].cells[index].paragraphs[0].add_run(value).bold = True

        table.rows[1].cells[0].text = "{}".format(report.sentences_count)
        table.rows[1].cells[1].text = "{} / {}".format(report.positive, report.negative)
        table.rows[1].cells[2].text = "{}".format(report.summary)

    @staticmethod
    def __add_sentences_reports_heading_row(table):
        row = table.add_row()
        row.cells[1].merge(row.cells[-1])
        for index, value in enumerate(["#", "Sentence"]):
            row.cells[index].paragraphs[0].add_run(value).bold = True

    @staticmethod
    def __add_sentence_report_row(table, report):
        row = table.add_row()
        row.cells[1].merge(row.cells[-1])
        row.cells[0].paragraphs[0].add_run("{}".format(report.index + 1))
        row.cells[1].paragraphs[0].add_run("{}".format(report.sentence)).italic = True
        return row

    @staticmethod
    def __add_keyword_reports_heading_row(table):
        row = table.add_row()
        for index, value in enumerate(["Category", "Keyword", "Evaluator", "Evaluation", "Inverted", "Intensifier"]):
            runner = row.cells[index + 1].paragraphs[0].add_run(value)
            runner.bold = True

    @staticmethod
    def __add_keyword_report_row(table, report):
        row = table.add_row()
        row.cells[1].text = "{}".format(report.category)
        row.cells[2].text = "{}".format(report.content)
        row.cells[3].text = report.evaluator if report.evaluator else '-'
        row.cells[4].text = "{} / {}".format(report.positive, report.negative)
        row.cells[5].text = 'Yes' if report.inverted else 'No'
        row.cells[6].text = report.intensifier if report.intensifier else '-'
        return row
