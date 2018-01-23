import settings

from docx import Document


class WordProcessor:

    @staticmethod
    def generate(report):
        filename = WordProcessor.__get_filename(report)

        document = Document()
        document.add_heading('Text report for "{}"'.format(report.excerpt), 1)

        document.add_heading('Input', 2)
        document.add_paragraph('{}'.format(report.content))

        document.add_heading('Summary', 2)
        summary_table = document.add_table(rows=2, cols=3)
        summary_table.style = 'Table Grid'
        WordProcessor.__add_summary_table_rows(summary_table, report)

        document.add_heading('Details', 2)
        details_table = document.add_table(rows=0, cols=6)
        details_table.style = 'Table Grid'
        WordProcessor.__add_sentences_reports_heading_row(details_table)
        for sentence_report in report.reports:
            first_row = WordProcessor.__add_sentence_report_row(details_table, sentence_report)
            WordProcessor.__add_keyword_reports_heading_row(details_table)
            for keyword_report in sentence_report.reports:
                last_row = WordProcessor.__add_keyword_report_row(details_table, keyword_report)
            first_row.cells[0].merge(last_row.cells[0])

        document.save('{0}/{1}.docx'.format(settings.EXPORT_PATH, filename))

    @staticmethod
    def __get_filename(report):
        return 'report'

    @staticmethod
    def __add_summary_table_rows(table, report):
        table.rows[0].cells[0].text = "Sentences"
        table.rows[0].cells[1].text = "Evaluation"
        table.rows[0].cells[2].text = "Summary"
        table.rows[1].cells[0].text = "{}".format(report.sentences_count)
        table.rows[1].cells[1].text = "{} / {}".format(report.positive, report.negative)
        table.rows[1].cells[2].text = "{}".format(report.summary)

    @staticmethod
    def __add_sentences_reports_heading_row(table):
        row = table.add_row()
        row.cells[1].merge(row.cells[-1])
        row.cells[0].text = "#"
        row.cells[1].text = "Sentence"

    @staticmethod
    def __add_sentence_report_row(table, report):
        row = table.add_row()
        row.cells[1].merge(row.cells[-1])
        row.cells[0].text = "{}".format(report.index + 1)
        row.cells[1].text = "{}".format(report.sentence)
        return row

    @staticmethod
    def __add_keyword_reports_heading_row(table):
        row = table.add_row()
        row.cells[1].text = "Category"
        row.cells[2].text = "Keyword"
        row.cells[3].text = "Evaluation"
        row.cells[4].text = "Inverted"
        row.cells[5].text = "Intensifier"

    @staticmethod
    def __add_keyword_report_row(table, report):
        row = table.add_row()
        row.cells[1].text = "{}".format(report.category)
        row.cells[2].text = "{}".format(report.content)
        row.cells[3].text = "{} / {}".format(report.positive, report.negative)
        row.cells[4].text = 'Yes' if report.inverted else 'No'
        row.cells[5].text = report.intensifier if report.intensifier else '-'
        return row
